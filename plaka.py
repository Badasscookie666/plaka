# preizo.py
from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Emu # Für die genaue Größenanpassung des Logos

import io
import os
import re

# Programmname von Plaka in Preizo geändert
preizo = Flask(__name__)

# --- Konfiguration für Logos ---
# Pfade zu den Logodateien im Ordner 'static'
KEMPER_LOGO_PATH = os.path.join(preizo.root_path, 'static', 'kemper_markt_logo.png')
AKTION_LOGO_PATH = os.path.join(preizo.root_path, 'static', 'sonder_angebot_logo.png')
BIO_LOGO_PATH = os.path.join(preizo.root_path, 'static', 'bio_logo.png')

# Definiere die Größen der Logos
KEMPER_LOGO_WIDTH = Inches(1.2)
AKTION_LOGO_WIDTH = Inches(1.35)
BIO_LOGO_WIDTH = Inches(1.6)

# Globale Schriftgrößen und Farben definiert
MAIN_FONT_SIZE = Pt(110)
SUB_FONT_SIZE = Pt(50)
DEPARTMENT_FONT_SIZE = Pt(35)
PRICE_FONT_SIZE = Pt(150)
DEPOSIT_FONT_SIZE = Pt(45)
RED = RGBColor(0xFF, 0x00, 0x00)

def set_spacing(paragraph, space_before, space_after):
    """Setzt den Abstand vor und nach einem Absatz."""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(space_before)
    paragraph_format.space_after = Pt(space_after)

def add_tight_paragraph(document, text, font_size, bold=False, font_color=None, spacing_after=Pt(0)):
    """Fügt einen Absatz mit angepasstem Zeilenabstand hinzu."""
    p = document.add_paragraph(text)
    run = p.runs[0]
    run.font.size = font_size
    run.font.name = 'Inter'
    run.font.bold = bold
    if font_color:
        run.font.color.rgb = font_color
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, 0, spacing_after)

def add_inline_image(paragraph, image_path, width=None):
    """Fügt ein Bild in einen Absatz ein und passt seine Größe an."""
    if os.path.exists(image_path):
        run = paragraph.add_run()
        run.add_picture(image_path, width=width)
        return True
    return False

def add_horizontal_line(document, width_inches=6.0):
    """Fügt eine horizontale Linie hinzu."""
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(15)
    p.paragraph_format.space_after = Pt(15)
    
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # Dicke der Linie
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def is_german_decimal(price_str):
    """Prüft, ob der String ein deutsches Dezimalformat hat (Komma statt Punkt)."""
    return re.match(r'^\d+,\d{2}$', price_str)

def clean_price(price_str):
    """Bereinigt den Preis-String und konvertiert ihn in einen Float."""
    price_str = price_str.strip().replace('€', '').replace(',', '.')
    return float(price_str)


# --- Hauptroute für das Formular ---
@preizo.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        # Erstelle ein neues Dokument
        document = Document()
        
        # Konfiguriere die Seite (Ränder)
        section = document.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
        # Leere Seitenränder
        document.sections[0].header_distance = Inches(0)
        document.sections[0].footer_distance = Inches(0)

        # Holen Sie sich die Formulardaten
        department = request.form.get('department', 'Abteilung nicht definiert')
        product_type = request.form.get('product_type', 'Normalpreis')
        manufacturer = request.form.get('manufacturer', '')
        product_name = request.form.get('product_name', 'Produktname nicht definiert')
        has_variety = 'has_variety' in request.form
        quantity = request.form.get('quantity', '')
        unit = request.form.get('unit', '')
        price_str = request.form.get('price', '')
        deposit_str = request.form.get('deposit', '0')
        packaging_type = request.form.get('packaging_type', '')
        is_gebinde = 'gebinde_check' in request.form
        gebinde_size_str = request.form.get('gebinde_size', '')
        inhalt_ml_str = request.form.get('inhalt_ml', '')
        
        should_be_red = product_type == 'Aktion'

        # Konvertiere Preis, Pfand und Gebinde-Daten
        try:
            price = clean_price(price_str)
        except (ValueError, TypeError):
            price = 0.0

        try:
            deposit = clean_price(deposit_str)
        except (ValueError, TypeError):
            deposit = 0.0

        try:
            gebinde_size = int(gebinde_size_str)
        except (ValueError, TypeError):
            gebinde_size = 0

        try:
            inhalt_ml = int(inhalt_ml_str)
        except (ValueError, TypeError):
            inhalt_ml = 0

        # --- Logo und Kopfzeile ---
        header_p = document.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Füge das Kemper-Logo hinzu
        kemper_logo_added = add_inline_image(header_p, KEMPER_LOGO_PATH, width=KEMPER_LOGO_WIDTH)

        if kemper_logo_added:
            kemper_run = header_p.add_run()
            kemper_run.add_text(' ') # Fügt einen Abstand hinzu
        
        # Füge das Bio- oder Aktionslogo hinzu, falls zutreffend
        if product_type == 'Bio':
            add_inline_image(header_p, BIO_LOGO_PATH, width=BIO_LOGO_WIDTH)
        elif product_type == 'Aktion':
            add_inline_image(header_p, AKTION_LOGO_PATH, width=AKTION_LOGO_WIDTH)

        set_spacing(header_p, 0, 0)
        
        # Füge eine horizontale Linie hinzu
        add_horizontal_line(document)

        # --- Abteilungsname ---
        add_tight_paragraph(document, department.upper(), DEPARTMENT_FONT_SIZE, bold=True, spacing_after=Pt(10))

        # --- Hersteller/Marke ---
        if manufacturer:
            add_tight_paragraph(document, manufacturer.upper(), SUB_FONT_SIZE, bold=False, spacing_after=Pt(20))

        # --- Produktname ---
        add_tight_paragraph(document, product_name.upper(), MAIN_FONT_SIZE, bold=True)
        if has_variety:
            add_tight_paragraph(document, "VERSCHIEDENE SORTEN", Pt(30), bold=True, font_color=RED if should_be_red else None)

        # --- Preis ---
        add_tight_paragraph(document, f"{price:,.2f}€".replace('.', ','), PRICE_FONT_SIZE, bold=True, font_color=RED if should_be_red else None, spacing_after=Pt(0))

        # --- Pfand ---
        if deposit > 0:
            add_tight_paragraph(document, f"Zzgl.: {deposit:,.2f}€ Pfand".replace('.', ','), DEPOSIT_FONT_SIZE, bold=True, font_color=RED if should_be_red else None)

        # --- Gebinde oder Einzelpreis pro Einheit ---
        if is_gebinde and gebinde_size > 0 and inhalt_ml > 0:
            total_volume_ml = gebinde_size * inhalt_ml
            if total_volume_ml > 0:
                price_per_liter_gebinde = (price / total_volume_ml) * 1000
                price_per_item_gebinde = price / gebinde_size

                gebinde_price_text = f"1 St.={price_per_item_gebinde:,.2f}€ / 1L={price_per_liter_gebinde:,.2f}€".replace('.', ',')
                add_tight_paragraph(document, gebinde_price_text, font_size=DEPOSIT_FONT_SIZE, bold=True, font_color=RED if should_be_red else None)
        else:
            # Berechnung für den Preis pro Einheit
            price_per_unit_text = ''
            quantity_float = float(quantity.replace(',', '.')) if quantity else 0
            if quantity_float > 0:
                if unit.lower() in ["g", "kg"]:
                    if unit.lower() == "g":
                        price_per_unit = (price / quantity_float) * 1000
                    else: # kg
                        price_per_unit = price / quantity_float
                    price_per_unit_text = f"1kg={price_per_unit:,.2f}€".replace('.', ',')
                elif unit.lower() in ["ml", "l"]:
                    if unit.lower() == "ml":
                        price_per_unit = (price / quantity_float) * 1000
                    else: # l
                        price_per_unit = price / quantity_float
                    price_per_unit_text = f"1L={price_per_unit:,.2f}€".replace('.', ',')
            
            if price_per_unit_text:
                add_tight_paragraph(document, price_per_unit_text, DEPOSIT_FONT_SIZE, bold=True, font_color=RED if should_be_red else None)

        # --- Packungsart ---
        if packaging_type:
             p = document.add_paragraph(packaging_type.upper())
             p.alignment = WD_ALIGN_PARAGRAPH.CENTER
             run = p.runs[0]
             run.font.size = Pt(DEPOSIT_FONT_SIZE)
             run.font.bold = True
             run.font.underline = True

        # --- Speichern oder Drucken ---
        output = io.BytesIO()
        document.save(output)
        output.seek(0)
        
        # Sende das generierte Dokument zum Download
        return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name='preisblatt.docx')

    return render_template('index.html')

if __name__ == '__main__':
    preizo.run(debug=True)
