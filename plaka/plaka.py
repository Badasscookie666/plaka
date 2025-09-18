# application.py
from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import io
import os

application = Flask(__name__)

# --- Configuration for logos ---
KEMPER_LOGO_PATH = os.path.join(application.root_path, 'static', 'kemper_markt_logo.png')
AKTION_LOGO_PATH = os.path.join(application.root_path, 'static', 'sonder_angebot_logo.png')
BIO_LOGO_PATH = os.path.join(application.root_path, 'static', 'bio_logo.png')

# Define consistent logo widths - increased by approx 10%
CONSISTENT_LOGO_WIDTH = Inches(1.1)
AKTION_LOGO_WIDTH = Inches(1.32)
BIO_LOGO_CUSTOM_WIDTH = Inches(1.5) # New variable for bigger BIO logo

# --- Route for the main form ---
@application.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        department = request.form['department']
        product_type = request.form['product_type']
        manufacturer = request.form.get('manufacturer', '').strip()
        product_name = request.form['product_name'].strip()
        has_varieties = request.form.get('has_varieties') == 'true'
        additional_info = request.form.get('additional_info', '').strip()

        quantity_per_pack_str = request.form['quantity_per_pack'].strip()
        unit = request.form['unit'].strip()
        price = float(request.form['price'])

        deposit = float(request.form.get('deposit', 0.0))
        packaging_type = request.form.get('packaging_type', '').strip()
        is_bio = request.form.get('is_bio') == 'true'
        
        price_category = request.form.get('price_category', '').strip()
        wiege_number = request.form.get('wiege_number', '').strip()

        # Gebinde specific fields
        is_gebinde = request.form.get('is_gebinde') == 'true'
        gebinde_size_str = request.form.get('gebinde_size', '').strip()
        inhalt_ml_str = request.form.get('inhalt_ml', '').strip()
        
        gebinde_size = 0
        inhalt_ml = 0
        try:
            if gebinde_size_str:
                gebinde_size = int(gebinde_size_str)
            if inhalt_ml_str:
                inhalt_ml = float(inhalt_ml_str)
        except ValueError:
            pass


        # Handle calculation of price per unit text
        price_per_unit_text = ""
        try:
            if unit.lower() in ["g", "kg", "ml", "l"]:
                calc_quantity = 0.0
                if '/' in quantity_per_pack_str:
                    quantities = [float(q) for q in quantity_per_pack_str.split('/')]
                    prices_per_unit_calc = []
                    for q in quantities:
                        if q == 0:
                            prices_per_unit_calc.append(0.0)
                            continue
                        if unit.lower() == "g":
                            prices_per_unit_calc.append((price / q) * 1000)
                        elif unit.lower() == "kg":
                            prices_per_unit_calc.append(price / q)
                        elif unit.lower() == "ml":
                            prices_per_unit_calc.append((price / q) * 1000)
                        elif unit.lower() == "l":
                            prices_per_unit_calc.append(price / q)

                    formatted_prices = [f"{p:,.2f}".replace('.', ',') for p in prices_per_unit_calc]
                    price_per_unit_text = "/".join(formatted_prices)
                    if unit.lower() in ["g", "kg"]:
                        price_per_unit_text = f"1kg={price_per_unit_text}€"
                    elif unit.lower() in ["ml", "l"]:
                        price_per_unit_text = f"1L={price_per_unit_text}€"
                else: # Single quantity
                    calc_quantity = float(quantity_per_pack_str)
                    if calc_quantity > 0:
                        if unit.lower() == "g":
                            price_per_kilo = (price / calc_quantity) * 1000
                            price_per_unit_text = f"1kg={price_per_kilo:,.2f}€".replace('.', ',')
                        elif unit.lower() == "kg":
                            price_per_kilo = price / calc_quantity
                            price_per_unit_text = f"1kg={price_per_kilo:,.2f}€".replace('.', ',')
                        elif unit.lower() == "ml":
                            price_per_liter = (price / calc_quantity) * 1000
                            price_per_unit_text = f"1L={price_per_liter:,.2f}€".replace('.', ',')
                        elif unit.lower() == "l":
                            price_per_liter = price / calc_quantity
                            price_per_unit_text = f"1L={price_per_liter:,.2f}€".replace('.', ',')
            # For other units like 'Stück', 'Tüte', 'Schale', 'Packung', 'Träger', no base price is calculated.
        except ValueError:
            pass

        document = generate_document(
            department,
            product_type,
            manufacturer,
            product_name,
            has_varieties,
            additional_info,
            quantity_per_pack_str,
            unit,
            price,
            deposit,
            packaging_type,
            price_per_unit_text,
            is_bio,
            price_category,
            wiege_number,
            is_gebinde,
            gebinde_size,
            inhalt_ml
        )

        doc_io = io.BytesIO()
        document.save(doc_io)
        doc_io.seek(0)

        filename_parts = []
        if manufacturer:
            filename_parts.append(manufacturer)
        filename_parts.append(product_name)
        filename_parts.append(f"{quantity_per_pack_str}{unit}")

        base_name = " ".join(filter(None, filename_parts)).strip()
        base_name = base_name.replace('/', '-')

        filename = f"{base_name}.docx"

        return send_file(doc_io, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    return render_template('index.html')

# --- Helper function to add a paragraph with specific style and tight spacing ---
def add_tight_paragraph(doc, text, font_name='Calibri', font_size=None, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=Pt(1), underline=False, font_color=None, highlight_text=False):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)

    run.font.name = font_name
    if font_size:
        run.font.size = Pt(font_size)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if underline or highlight_text: # This condition now handles both
        run.underline = True
    if font_color:
        run.font.color.rgb = font_color

    paragraph.alignment = alignment

    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = spacing_after

    if highlight_text:
        run.font.highlight_color = WD_COLOR_INDEX.GRAY_25

    return paragraph

# --- Function to generate the Word document ---
def generate_document(department, product_type, manufacturer, product_name,
                    has_varieties, additional_info, quantity_per_pack_str, unit, price, deposit,
                    packaging_type, price_per_unit_text, is_bio, price_category, wiege_number,
                    is_gebinde, gebinde_size, inhalt_ml): # Added new parameters
    document = Document()
    section = document.sections[0]

    # Set very small margins to maximize content area
    section.left_margin = Inches(0.2)
    section.right_margin = Inches(0.2)
    section.top_margin = Inches(0.2)
    section.bottom_margin = Inches(0.2)

    # Set page size (approx. 4x6 inches seems reasonable for price tags)
    section.page_width = Inches(4)
    section.page_height = Inches(6) # Total page height is 6 inches = 432 Pt

    # Set default font for the document (applied to 'Normal' style)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11) # Default font size

    # Define RED color
    RED = RGBColor(0xFF, 0x00, 0x00)

    # Determine if prices should be red
    should_be_red = (product_type == 'Aktion')

    # --- Logo Selection and Insertion based on Hierarchy ---
    logo_to_insert_path = None
    logo_to_insert_width = CONSISTENT_LOGO_WIDTH # Default width

    if department == 'Obst&Gemüse' and is_bio and os.path.exists(BIO_LOGO_PATH):
        logo_to_insert_path = BIO_LOGO_PATH
        logo_to_insert_width = BIO_LOGO_CUSTOM_WIDTH # Use custom width for BIO logo
    elif product_type == 'Aktion' and os.path.exists(AKTION_LOGO_PATH):
        logo_to_insert_path = AKTION_LOGO_PATH
        logo_to_insert_width = AKTION_LOGO_WIDTH
    elif os.path.exists(KEMPER_LOGO_PATH): # Default if neither Bio nor Aktion
        logo_to_insert_path = KEMPER_LOGO_PATH
        logo_to_insert_width = CONSISTENT_LOGO_WIDTH
    
    if logo_to_insert_path:
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(logo_to_insert_path, width=logo_to_insert_width)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(0) # No space after the logo

    # Add an empty line between logo and first text element ONLY for Obst&Gemüse
    if department == 'Obst&Gemüse':
        add_tight_paragraph(document, "", spacing_after=Pt(12), font_size=38)

    # Define common font sizes for consistency across departments where applicable
    PRODUCT_NAME_FONT_SIZE = 26 # Consistent product name font size for all categories

    if department == 'Obst&Gemüse':
        MANUFACTURER_FONT_SIZE = 30 # Specific to Obst&Gemüse
        ADDITIONAL_INFO_FONT_SIZE = 18
        QTY_UNIT_FONT_SIZE = 20
        PRICE_FONT_SIZE = 65
        PRICE_UNIT_FONT_SIZE = 16
        WIEGE_NR_FONT_SIZE = 28

        if manufacturer:
            add_tight_paragraph(document, f"{manufacturer}", font_size=MANUFACTURER_FONT_SIZE, bold=True, spacing_after=Pt(0)) # Spacing adjusted
        else: # Add empty line placeholder if no manufacturer
            add_tight_paragraph(document, "", font_size=MANUFACTURER_FONT_SIZE, bold=True, spacing_after=Pt(0)) # Spacing adjusted
        add_tight_paragraph(document, f"{product_name}", font_size=PRODUCT_NAME_FONT_SIZE, bold=True, spacing_after=Pt(0)) # Spacing adjusted
        
        if additional_info:
            add_tight_paragraph(document, f"{additional_info}", font_size=ADDITIONAL_INFO_FONT_SIZE)
        
        if price_category == 'WIEGE_NR' and wiege_number:
            add_tight_paragraph(document, f"Wiege Nr. {wiege_number}", font_size=WIEGE_NR_FONT_SIZE, bold=True, highlight_text=True)
        else:
            container_word_obsgem = "Packung"
            if "tüte" in unit.lower():
                container_word_obsgem = "Tüte"
            elif "schale" in unit.lower():
                container_word_obsgem = "Schale"
            elif "stück" in unit.lower():
                container_word_obsgem = "Stück"
            elif "flasche" in unit.lower():
                container_word_obsgem = "Flasche"

            je_line_text_obsgem = f"Je {quantity_per_pack_str}{unit} {container_word_obsgem}"
            add_tight_paragraph(document, je_line_text_obsgem, font_size=QTY_UNIT_FONT_SIZE, highlight_text=True)

        add_tight_paragraph(document, f"{price:,.2f}€".replace('.', ','), font_name='Times New Roman', font_size=PRICE_FONT_SIZE, bold=True, font_color=RED if should_be_red else None, spacing_after=Pt(0)) # Spacing adjusted
        if price_per_unit_text:
            add_tight_paragraph(document, price_per_unit_text, font_name='Times New Roman', font_size=PRICE_UNIT_FONT_SIZE, font_color=RED if should_be_red else None)

    elif department == 'Trocken Sortiment':
        MANUFACTURER_FONT_SIZE = 36 # Specific to Trocken Sortiment
        VARIETIES_FONT_SIZE = 20
        QTY_UNIT_FONT_SIZE = 22
        PRICE_FONT_SIZE = 70
        PRICE_UNIT_FONT_SIZE = 18

        if manufacturer:
            add_tight_paragraph(document, f"{manufacturer}", font_name='Calibri', font_size=MANUFACTURER_FONT_SIZE, bold=True, spacing_after=Pt(0)) # Spacing adjusted
        else: # Add empty line placeholder if no manufacturer
            add_tight_paragraph(document, "", font_name='Calibri', font_size=MANUFACTURER_FONT_SIZE, bold=True, spacing_after=Pt(0)) # Spacing adjusted
        
        add_tight_paragraph(document, f"{product_name}", font_name='Calibri', font_size=PRODUCT_NAME_FONT_SIZE, bold=True, spacing_after=Pt(0)) # Spacing adjusted
        
        if has_varieties:
            add_tight_paragraph(document, "Verschiedene Sorten", font_name='Calibri', font_size=VARIETIES_FONT_SIZE)
            
        container_word = "Packung"
        if "glas" in quantity_per_pack_str.lower() or "glas" in unit.lower():
            container_word = "Glas"
        elif "dose" in quantity_per_pack_str.lower() or "dose" in unit.lower():
            container_word = "Dose"
        elif unit.lower() == "stück":
            container_word = "Stück"
        elif unit.lower() == "tüte":
            container_word = "Tüte"
        elif unit.lower() == "schale":
            container_word = "Schale"
        elif unit.lower() == "träger":
            container_word = "Träger"
        elif unit.lower() == "kg" and not any(word in product_name.lower() for word in ["beutel", "packung", "tüte", "schale"]):
            container_word = "Packung" 

        je_line_text = f"Je {quantity_per_pack_str}{unit} {container_word}"
        add_tight_paragraph(document, je_line_text, font_name='Calibri', font_size=QTY_UNIT_FONT_SIZE, highlight_text=True)

        add_tight_paragraph(document, f"{price:,.2f}€".replace('.', ','), font_name='Times New Roman', font_size=PRICE_FONT_SIZE, bold=True, font_color=RED if should_be_red else None, spacing_after=Pt(0)) # Spacing adjusted
        
        if price_per_unit_text:
            add_tight_paragraph(document, price_per_unit_text, font_name='Times New Roman', font_size=PRICE_UNIT_FONT_SIZE, font_color=RED if should_be_red else None)

    elif department == 'Getränke':
        MANUFACTURER_FONT_SIZE = 36 # Specific to Getränke
        VARIETIES_FONT_SIZE = 20
        QTY_UNIT_FONT_SIZE = 22
        PRICE_FONT_SIZE = 70
        PRICE_UNIT_FONT_SIZE = 18
        DEPOSIT_FONT_SIZE = 14
        PACKAGING_TYPE_FONT_SIZE = 30

        if manufacturer:
            add_tight_paragraph(document, f"{manufacturer}", font_size=MANUFACTURER_FONT_SIZE, bold=True, spacing_after=Pt(0)) # Spacing adjusted
        else: # Add empty line placeholder if no manufacturer
            add_tight_paragraph(document, "", font_size=MANUFACTURER_FONT_SIZE, bold=True, spacing_after=Pt(0)) # Spacing adjusted
        
        add_tight_paragraph(document, f"{product_name}", font_size=PRODUCT_NAME_FONT_SIZE, bold=True, spacing_after=Pt(0)) # Spacing adjusted

        if has_varieties:
            add_tight_paragraph(document, "Verschiedene Sorten", font_size=VARIETIES_FONT_SIZE)

        # The je_line_text for Getränke now correctly only uses quantity_per_pack_str and unit
        je_line_text = f"Je {quantity_per_pack_str}{unit}"
        add_tight_paragraph(document, je_line_text, font_size=QTY_UNIT_FONT_SIZE, highlight_text=True) # Apply highlight to Getränke je_line_text

        add_tight_paragraph(document, f"{price:,.2f}€".replace('.', ','), font_name='Times New Roman', font_size=PRICE_FONT_SIZE, bold=True, font_color=RED if should_be_red else None, spacing_after=Pt(0)) # Spacing adjusted
        
        if deposit > 0:
            add_tight_paragraph(document, f"Zzgl.: {deposit:,.2f}€ Pfand".replace('.', ','), font_size=DEPOSIT_FONT_SIZE, bold=True, font_color=RED if should_be_red else None)
        
        # Gebinde Price Per Unit/Liter calculation and display
        if is_gebinde and gebinde_size > 0 and inhalt_ml > 0:
            total_volume_ml = gebinde_size * inhalt_ml
            if total_volume_ml > 0:
                price_per_liter_gebinde = (price / total_volume_ml) * 1000
                price_per_item_gebinde = price / gebinde_size

                gebinde_price_text = f"1 St.={price_per_item_gebinde:,.2f}€ / 1L={price_per_liter_gebinde:,.2f}€".replace('.', ',')
                add_tight_paragraph(document, gebinde_price_text, font_size=DEPOSIT_FONT_SIZE, bold=True, font_color=RED if should_be_red else None)


        if price_per_unit_text and not is_gebinde: # Display general price_per_unit_text only if not Gebinde
            add_tight_paragraph(document, price_per_unit_text, font_name='Times New Roman', font_size=PRICE_UNIT_FONT_SIZE, font_color=RED if should_be_red else None)
        
        if packaging_type:
            add_tight_paragraph(document, packaging_type.upper(), font_size=PACKAGING_TYPE_FONT_SIZE, bold=True, underline=True)

    return document

if __name__ == '__main__':
    application.run(debug=True)
